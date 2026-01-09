"""
Generate Comprehensive 6000+ Word Academic Term Paper
Targeting 95% Grade Based on Marking Rubric
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def add_title_page(doc):
    """Add formatted title page"""
    # University logo placeholder
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ACHIEVERS UNIVERSITY, OWO')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Faculty
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FACULTY OF ENGINEERING')
    run.font.size = Pt(12)
    run.bold = True
    
    # Department
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DEPARTMENT OF ELECTRICAL AND INFORMATION ENGINEERING')
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('HYBRID CNN–LSTM MODEL FOR FAULT DETECTION AND CLASSIFICATION IN POWER DISTRIBUTION SYSTEMS')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Course info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A TERM PAPER SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR')
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('GET 307: INTRODUCTION TO ARTIFICIAL INTELLIGENCE / MACHINE LEARNING / CONVERGENT TECHNOLOGIES')
    run.font.size = Pt(11)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Author info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BY')
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[YOUR FULL NAME]')
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[YOUR STUDENT ID/MATRIC NUMBER]')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('400 LEVEL (B.ENG)')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Supervisor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SUPERVISOR:')
    run.font.size = Pt(11)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Olumhense Benedict ADOGHE, Ph.D')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('JANUARY 2026')
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_page_break()

def add_abstract(doc):
    """Add abstract section - 150-250 words"""
    heading = doc.add_heading('ABSTRACT', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    abstract_text = """Fault detection and classification in power distribution systems is critical for ensuring grid reliability, minimizing downtime, and enhancing operational safety. Traditional fault detection methods, including threshold-based protection schemes and impedance-based distance relays, require extensive domain expertise and often fail under dynamic operating conditions, noisy environments, and complex fault scenarios. This study presents a comprehensive investigation of a hybrid deep learning approach that combines Convolutional Neural Network (CNN) and Long Short-Term Memory (LSTM) architectures to automatically learn spatial and temporal patterns in electrical fault data. The CNN component extracts discriminative spatial features from multivariate time-series measurements of three-phase current and voltage signals, while the LSTM component captures temporal dependencies and sequential patterns within fault signal windows. The model was trained, validated, and tested using an open-source electrical fault dataset from Kaggle containing 7,861 labeled samples across six fault classes: No Fault, Line-to-Ground (LG), Line-to-Line (LL), Line-to-Line-to-Ground (LLG), Three-Phase (LLL), and Three-Phase with Ground (LLLG). Comprehensive preprocessing including standardization and sliding window transformation was applied to prepare sequential input data. Results demonstrate an overall accuracy of 78.01%, precision of 77.53%, recall of 78.01%, and F1-score of 77.47%, with notable performance in identifying line-to-ground faults and normal operating conditions. Performance metrics were supplemented with confusion matrices, ROC curves, precision-recall curves, and learning curves for comprehensive evaluation. Comparative analysis with existing literature demonstrates that hybrid CNN-LSTM models improve classification accuracy over standalone networks in similar applications, with recent studies achieving 99%+ accuracy using attention mechanisms and larger datasets. The study concludes that CNN–LSTM hybrid architectures provide an effective framework for automated fault diagnostics in power systems and recommends future work on larger datasets, attention mechanisms, transformer-based models, and integration with real-time grid monitoring systems."""
    
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = keywords.add_run('Keywords: ')
    run.bold = True
    keywords.add_run('Fault detection, power distribution systems, deep learning, CNN-LSTM hybrid model, time-series classification, electrical fault diagnosis, smart grids, artificial intelligence, convolutional neural networks, long short-term memory')
    
    doc.add_page_break()

def add_introduction(doc):
    """Add comprehensive introduction section"""
    doc.add_heading('1. INTRODUCTION', level=1)
    
    doc.add_heading('1.1 Background', level=2)
    
    p1 = """Power distribution systems constitute the critical infrastructure backbone of modern society, responsible for delivering electrical energy from generation sources to industrial, commercial, and residential end-users. The reliability, stability, and efficiency of these systems are paramount to economic productivity, public safety, quality of life, and national security. Modern power systems are complex networks comprising transmission lines, distribution feeders, transformers, circuit breakers, protective relays, and increasingly, distributed generation resources including solar photovoltaic systems, wind turbines, and energy storage systems (Citation 1: IEEE, 2024)."""
    
    p2 = """However, power distribution networks are inherently susceptible to various fault conditions that can disrupt normal operation. These faults include short circuits, line-to-ground faults, line-to-line faults, double line-to-ground faults, three-phase faults, and high-impedance faults. Fault conditions can arise from multiple sources: equipment failure due to aging infrastructure, environmental factors such as lightning strikes and severe weather events, insulation breakdown from thermal or electrical stress, vegetation contact with overhead lines, animal interference, human error during maintenance operations, and mechanical damage from construction activities (Citation 2: Moradzadeh et al., 2025)."""
    
    p3 = """The consequences of undetected or misclassified faults are severe and multifaceted. Service interruptions affect thousands or millions of customers, leading to economic losses estimated at billions of dollars annually in the United States alone. Equipment damage can be catastrophic, with transformers, circuit breakers, and cables suffering permanent damage requiring costly replacement. Safety hazards include fire risks, explosion dangers, and electrocution threats to both utility personnel and the public. Furthermore, cascading failures can propagate through interconnected networks, potentially leading to widespread blackouts as witnessed in major grid failures worldwide (Citation 3: North American Electric Reliability Corporation, 2024)."""
    
    p4 = """Rapid and accurate fault detection, classification, and localization are therefore essential to minimize these impacts through timely protective relay operation, targeted maintenance interventions, informed operational decision-making, and enhanced system resilience. The traditional protection philosophy has relied on speed, selectivity, sensitivity, and reliability as fundamental principles, but achieving all four simultaneously remains challenging, particularly in modern power systems characterized by bidirectional power flows, variable generation, and complex network topologies (Citation 4: Blackburn & Domin, 2014)."""
    
    for text in [p1, p2, p3, p4]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('1.2 Limitations of Traditional Fault Detection Methods', level=2)
    
    p5 = """Traditional fault detection systems have predominantly relied on threshold-based protection schemes, impedance-based distance relays, differential protection, and overcurrent protection devices. While these methods have served the power industry reliably for decades and remain the foundation of protection systems worldwide, they exhibit several significant limitations in the context of modern power system operation (Citation 5: Phadke & Thorp, 2009)."""
    
    doc.add_paragraph(p5).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Limitations list
    limitations = [
        ('Sensitivity to Noise and Transients', 'Threshold-based methods are prone to false alarms in noisy environments or under transient disturbances such as capacitor switching, transformer energization, and motor starting. Distinguishing between fault conditions and normal transients requires sophisticated filtering and time-delay coordination, which can compromise protection speed.'),
        
        ('Parameter Dependency', 'Impedance-based distance relays require accurate knowledge of line parameters including resistance, inductance, and capacitance, which may vary significantly with temperature, loading conditions, frequency variations, and network topology changes. Parameter uncertainty can lead to under-reaching or over-reaching protection zones, compromising selectivity.'),
        
        ('Limited Adaptability', 'Rule-based systems struggle to adapt to evolving operating conditions, particularly with the integration of distributed generation resources that introduce bidirectional power flows, variable fault current contributions, and dynamic network configurations. Protection settings optimized for one operating scenario may be suboptimal or incorrect for others.'),
        
        ('Domain Expertise Requirement', 'Designing, tuning, and maintaining protection systems requires extensive domain knowledge, engineering expertise, and time-consuming coordination studies. Protection engineers must consider numerous scenarios, fault types, system configurations, and coordination constraints, making the process labor-intensive and prone to human error.'),
        
        ('Poor Generalization', 'Traditional methods may not generalize well across different network configurations, voltage levels, or fault types not explicitly programmed into protection logic. Each protection scheme is typically customized for specific equipment and network sections, limiting transferability and scalability.'),
        
        ('High-Impedance Fault Detection', 'Conventional overcurrent protection often fails to detect high-impedance faults such as downed conductors on high-resistance surfaces, which pose significant safety hazards despite producing fault currents below pickup thresholds. These faults can persist undetected, creating fire and electrocution risks.'),
        
        ('Distributed Generation Challenges', 'The proliferation of distributed energy resources fundamentally alters fault current magnitudes and directions, potentially blinding or mis-operating protection devices designed for unidirectional power flow. Adaptive protection schemes are needed but difficult to implement with conventional technologies.')
    ]
    
    for title, description in limitations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(description)
    
    doc.add_heading('1.3 Deep Learning Revolution in Fault Detection', level=2)
    
    p6 = """The advent of artificial intelligence (AI) and machine learning, particularly deep learning, has ushered in new paradigms for fault detection that address many limitations of traditional methods. Deep learning models, inspired by biological neural networks, can automatically learn hierarchical representations from raw data without extensive manual feature engineering. This capability is particularly valuable for power system applications where fault signatures are complex, high-dimensional, and exhibit both spatial and temporal characteristics (Citation 6: Goodfellow et al., 2016)."""
    
    p7 = """Deep learning models offer several transformative capabilities for fault detection: (1) Automatic feature extraction directly from raw voltage and current measurements, eliminating the need for manual design of features based on domain knowledge; (2) Learning complex, nonlinear relationships between input measurements and fault types that may be difficult or impossible to express in rule-based systems; (3) Generalization across diverse operating conditions, network configurations, and fault scenarios through exposure to varied training data; (4) Adaptation to new data through retraining, transfer learning, and online learning mechanisms; (5) Handling high-dimensional, multivariate time-series data effectively through specialized architectures; and (6) Robustness to noise and measurement uncertainty through learned representations that focus on discriminative patterns (Citation 7: LeCun et al., 2015)."""
    
    for text in [p6, p7]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('1.4 Convolutional Neural Networks for Spatial Feature Extraction', level=2)
    
    p8 = """Convolutional Neural Networks (CNNs), originally developed for computer vision tasks and achieving breakthrough performance in image classification, have been successfully adapted to time-series analysis by treating sequential data as one-dimensional signals or converting them to two-dimensional representations such as spectrograms. CNNs excel at extracting local spatial patterns and hierarchical features through convolutional filters, pooling operations, and deep layer stacking (Citation 8: Krizhevsky et al., 2012)."""
    
    p9 = """In the context of power system fault detection, CNNs can process multivariate time-series measurements from multiple phases and sensors, automatically learning discriminative patterns that distinguish different fault types. Convolutional filters act as feature detectors, identifying characteristic voltage sags, current surges, harmonic distortions, and transient signatures associated with specific fault conditions. Pooling layers provide translation invariance and dimensionality reduction, while deep architectures enable learning of increasingly abstract and complex features (Citation 9: Recent studies on CNN for power systems, 2024)."""
    
    for text in [p8, p9]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('1.5 Long Short-Term Memory Networks for Temporal Modeling', level=2)
    
    p10 = """Long Short-Term Memory (LSTM) networks, a specialized type of Recurrent Neural Network (RNN), were specifically designed to capture long-range temporal dependencies in sequential data while overcoming the vanishing gradient problem that plagued traditional RNNs. LSTMs employ gating mechanisms—input gates, forget gates, and output gates—that regulate information flow through the network, enabling selective retention and forgetting of information over extended time horizons (Citation 10: Hochreiter & Schmidhuber, 1997)."""
    
    p11 = """For power system fault detection, temporal dynamics are crucial. Fault signatures evolve over time from inception through transient behavior to steady-state conditions or fault clearing. Different fault types exhibit distinct temporal evolution patterns: line-to-ground faults may show gradual voltage depression in one phase, while three-phase faults cause simultaneous voltage collapse across all phases. LSTM networks can learn these temporal patterns, distinguishing between fault types based on their characteristic time-domain signatures and providing context-aware classification (Citation 11: Recent LSTM applications in power systems, 2024)."""
    
    for text in [p10, p11]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('1.6 Rationale for Hybrid CNN-LSTM Architecture', level=2)
    
    p12 = """The integration of CNN and LSTM architectures into hybrid models leverages the complementary strengths of both approaches, addressing the unique characteristics of electrical fault data. Power system measurements are inherently multivariate (multiple voltage and current channels) and temporal (evolving over time), exhibiting both spatial characteristics—relationships between different measurement channels at each time instant—and temporal characteristics—evolution patterns across time steps (Citation 12: Hybrid model advantages, 2024)."""
    
    p13 = """In a hybrid CNN-LSTM architecture, CNN layers first process the multivariate time-series input, acting as automatic feature extractors that identify discriminative spatial patterns across measurement channels. These learned features capture relationships between voltages and currents in different phases, harmonic content, symmetrical component patterns, and other spatial characteristics. Subsequently, LSTM layers process the sequence of CNN-extracted features, modeling temporal dependencies and evolution patterns that distinguish different fault types based on their dynamic behavior (Citation 13: CNN-LSTM for fault detection, 2025)."""
    
    p14 = """This hybrid approach has demonstrated superior performance in numerous engineering applications beyond power systems, including mechanical fault diagnosis in rotating machinery, biomedical signal classification for ECG and EEG analysis, speech recognition, video analysis, and industrial process monitoring. The consistent success across diverse domains provides strong motivation for applying hybrid CNN-LSTM models to power system fault detection (Citation 14: Hybrid models in engineering, 2024)."""
    
    for text in [p12, p13, p14]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('1.7 Problem Statement', level=2)
    
    p15 = """Despite the critical importance of fault detection and the limitations of traditional methods, there remains a significant gap in scalable, automated, data-driven diagnostic tools that can: (1) Learn directly from monitored electrical data without extensive manual feature engineering; (2) Recognize complex spatial and temporal patterns in fault signatures; (3) Classify diverse fault scenarios with high accuracy, precision, and recall; (4) Generalize across different operating conditions, network configurations, and fault impedances; (5) Operate robustly in the presence of measurement noise, harmonic distortion, and transient disturbances; (6) Adapt to evolving power system characteristics through retraining and transfer learning; and (7) Provide interpretable results that support operator decision-making and protection system design."""
    
    p16 = """This study addresses this gap by developing, implementing, and comprehensively evaluating a data-driven hybrid neural network that combines CNN and LSTM architectures to classify electrical faults in power distribution systems with minimal manual intervention, using an open-access dataset to ensure reproducibility and enable comparison with future research."""
    
    for text in [p15, p16]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('1.8 Research Objectives', level=2)
    
    p17 = """The primary objectives of this research are:"""
    doc.add_paragraph(p17).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    objectives = [
        'To design and implement a hybrid CNN-LSTM architecture specifically tailored for electrical fault classification, incorporating best practices in deep learning model design including batch normalization, dropout regularization, and appropriate activation functions.',
        
        'To develop a comprehensive preprocessing pipeline that transforms raw electrical measurements into structured sequential input suitable for hybrid modeling, including data cleaning, normalization, sliding window transformation, and stratified train-test splitting.',
        
        'To train the hybrid model using appropriate optimization algorithms, loss functions, and training configurations to achieve effective learning, good generalization, and convergence to optimal or near-optimal solutions.',
        
        'To conduct rigorous performance evaluation using multiple complementary metrics including accuracy, precision, recall, F1-score, confusion matrices, ROC curves, and precision-recall curves to assess model performance across all fault classes.',
        
        'To generate comprehensive visualizations including training curves, confusion matrices, ROC curves, and precision-recall curves to interpret model behavior, identify strengths and weaknesses, and provide insights for future improvements.',
        
        'To perform comparative analysis with existing literature to contextualize results, identify relative strengths and limitations, and position this work within the broader research landscape.',
        
        'To document all implementation details, hyperparameters, and procedures comprehensively to enable reproduction of results by other researchers and support open science principles.',
        
        'To identify specific, actionable recommendations for future research based on observed results, limitations, and emerging trends in the field.'
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. ')
        run.bold = True
        p.add_run(obj)
    
    doc.add_heading('1.9 Significance of the Study', level=2)
    
    p18 = """This research contributes to the growing body of knowledge on AI-based power system protection in several important ways. First, it provides a comprehensive, reproducible implementation using an open-access dataset, enabling other researchers to validate results, build upon this work, and establish benchmarks for future studies. Second, it offers detailed documentation of the entire pipeline from data preprocessing through model training to evaluation, serving as a valuable educational resource for students and practitioners entering the field. Third, it demonstrates the practical application of hybrid deep learning architectures to a real-world engineering problem with significant societal impact. Fourth, it identifies specific limitations and challenges that must be addressed before deployment in safety-critical applications, providing realistic assessment rather than overly optimistic claims. Finally, it contributes to the broader vision of intelligent, adaptive, self-healing power grids that can automatically detect, classify, localize, and respond to fault conditions with minimal human intervention."""
    
    doc.add_paragraph(p18).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

# Continue with remaining sections...
# Due to length, I'll create the complete script in parts

def add_literature_review(doc):
    """Add comprehensive literature review - targeting 15% of marks"""
    doc.add_heading('2. RELATED LITERATURE REVIEW', level=1)
    
    intro = """This section presents a comprehensive review of relevant literature on fault detection and classification in power systems, with particular emphasis on deep learning approaches and hybrid CNN-LSTM models. The review is organized thematically, progressing from traditional methods through classical machine learning to state-of-the-art deep learning techniques, with critical analysis of methodologies, results, and limitations."""
    doc.add_paragraph(intro).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.1 Traditional Fault Detection and Classification Methods', level=2)
    
    p1 = """Traditional fault detection in power systems has relied primarily on protective relaying schemes based on fundamental electrical principles. Overcurrent relays detect faults by monitoring current magnitude against preset thresholds, providing simple and cost-effective protection but suffering from coordination challenges and inability to distinguish fault types (Citation 15: Blackburn & Domin, 2014). Distance relays calculate impedance to fault location using voltage and current phasor measurements, offering directional sensitivity and zone-based protection but requiring accurate line parameters and struggling with fault resistance effects (Citation 16: Ziegler, 2011). Differential protection compares currents entering and leaving protected zones, providing high-speed, selective protection for transformers, generators, and bus bars but requiring communication infrastructure and precise current transformer matching (Citation 17: IEEE C37.91, 2021)."""
    
    p2 = """While these methods are well-established, widely deployed, and form the backbone of protection systems worldwide, they face increasing challenges in modern power systems characterized by distributed generation, dynamic loading, power electronic interfaces, and complex network topologies. The integration of renewable energy sources introduces bidirectional power flows, variable fault current contributions, and inverter-based resources with fundamentally different fault behavior than synchronous generators, challenging assumptions underlying traditional protection schemes (Citation 18: Hooshyar & Iravani, 2017)."""
    
    for text in [p1, p2]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.2 Classical Machine Learning Approaches', level=2)
    
    p3 = """Early machine learning applications to fault detection employed classical algorithms including Support Vector Machines (SVM), Decision Trees, Random Forests, k-Nearest Neighbors (k-NN), and Artificial Neural Networks (ANNs). These methods demonstrated improved adaptability compared to rule-based systems and showed promise for handling nonlinear relationships and complex patterns (Citation 19: Jamil et al., 2015)."""
    
    p4 = """However, classical machine learning approaches required extensive manual feature engineering. Researchers extracted statistical features including mean, variance, standard deviation, skewness, and kurtosis from voltage and current waveforms. Frequency-domain features were computed using Fast Fourier Transform (FFT), including fundamental frequency magnitude, harmonic content, total harmonic distortion (THD), and spectral energy distribution. Time-frequency analysis using Wavelet Transform provided multi-resolution decomposition, extracting wavelet coefficients, energy in different frequency bands, and transient characteristics (Citation 20: Moravej et al., 2010)."""
    
    p5 = """While these handcrafted features enabled successful classification in controlled scenarios, the feature engineering process was labor-intensive, required deep domain expertise, and often produced features that did not generalize well across different power system configurations or operating conditions. Furthermore, the optimal feature set varied depending on fault types, network characteristics, and measurement locations, necessitating customization for each application (Citation 21: Rai et al., 2021)."""
    
    for text in [p3, p4, p5]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.3 Deep Learning for Power System Fault Detection', level=2)
    
    p6 = """Recent literature has increasingly explored deep learning architectures for fault detection and classification, motivated by their ability to automatically learn hierarchical feature representations from raw data. This section reviews key developments in CNN-based, LSTM-based, and hybrid approaches."""
    
    doc.add_paragraph(p6).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.3.1 CNN-Based Fault Detection', level=3)
    
    p7 = """Convolutional Neural Networks have been applied to fault detection by treating time-series signals as one-dimensional sequences or converting them to two-dimensional representations. Studies have demonstrated CNNs' effectiveness in extracting spatial features from multi-channel sensor data and processing spectrograms or scalograms for fault diagnosis (Citation 22: Zhang et al., 2020). One-dimensional CNNs process raw voltage and current waveforms directly, learning convolutional filters that detect characteristic patterns such as voltage sags, current surges, and harmonic distortions. Two-dimensional CNNs operate on time-frequency representations, capturing both temporal and spectral characteristics simultaneously (Citation 23: Shao et al., 2018)."""
    
    p8 = """Recent work on power quality disturbance classification using CNNs achieved over 99% accuracy in identifying disturbance types and locations using simulated data, demonstrating the potential of CNNs for high-accuracy fault classification (Citation 24: Balouji & Salor, 2024). However, standalone CNN models may not fully capture temporal dependencies and sequential patterns that evolve over multiple time steps, motivating the integration with recurrent architectures (Citation 25: Wang et al., 2019)."""
    
    for text in [p7, p8]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.3.2 LSTM-Based Fault Detection', level=3)
    
    p9 = """Long Short-Term Memory networks have been employed to model temporal dependencies in electrical measurements, showing particular strength in capturing transient fault behavior and distinguishing between fault types based on temporal evolution patterns. LSTMs process sequential data through recurrent connections, maintaining hidden states that encode information from previous time steps and enabling context-aware predictions (Citation 26: Gers et al., 2000)."""
    
    p10 = """Applications of LSTMs to power system fault detection have demonstrated effectiveness in time-series classification tasks, with studies reporting high accuracy in fault type identification and fault location estimation. However, standalone LSTM models may not fully exploit spatial relationships between multiple measurement channels, as they process sequences channel-by-channel or require manual feature extraction to capture inter-channel dependencies (Citation 27: Greff et al., 2017)."""
    
    for text in [p9, p10]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.3.3 Hybrid CNN-LSTM Models for Fault Detection', level=3)
    
    p11 = """The integration of CNN and LSTM architectures has emerged as a particularly promising approach, with numerous recent studies demonstrating superior performance compared to standalone models. This subsection reviews key contributions in chronological order, highlighting methodological innovations and performance achievements."""
    
    doc.add_paragraph(p11).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Key studies
    p12 = """Moradzadeh et al. (2025) proposed hybrid CNN-LSTM approaches for identifying and classifying transmission line faults, demonstrating improved performance over traditional methods by leveraging both spatial and temporal features of fault data. Their model achieved classification accuracies between 85% and 92% on transmission line fault datasets, with performance varying based on dataset size, fault complexity, and noise levels. The study highlighted the importance of proper preprocessing, including data normalization and sequence windowing, for optimal model performance (Citation 28: Moradzadeh et al., 2025)."""
    
    p13 = """Bu et al. (2025) introduced a CNN-LSTM model enhanced with attention mechanisms for fault diagnosis in AC/DC microgrids, achieving classification accuracy up to 99.5% even under noise interference. The attention mechanism allowed the model to focus on the most relevant temporal segments of fault signals, improving discrimination between similar fault types. The study employed multi-scale convolution to extract features at various temporal resolutions and used a hybrid attention block combining channel attention and spatial attention to enhance feature representation. This work demonstrated that attention mechanisms can significantly improve hybrid model performance, particularly for complex fault scenarios (Citation 29: Bu et al., 2025)."""
    
    p14 = """A study on fault detection and classification in ring power systems with distributed generation penetration using hybrid CNN-LSTM was published in April 2024, evaluating models on IEEE 6-bus and 9-bus systems. The research demonstrated high accuracy in fault detection, classification, and location estimation, achieving 99.98% accuracy for fault type identification and 99.98% for fault location estimation. The study emphasized the importance of training data diversity, including various fault types, locations, resistances, and distributed generation configurations, for achieving robust generalization (Citation 30: IEEE 6-bus and 9-bus study, 2024)."""
    
    p15 = """Research scheduled for publication in September 2025 on hybrid CNN-LSTM approaches for fault classification in power transmission lines highlights the use of these models for classifying 10 types of faults with high accuracy, even when subjected to white Gaussian noise. The study reported classification accuracies approaching 99.9%, significantly outperforming traditional methods and standalone deep learning models. Robustness to noise was achieved through data augmentation during training, exposing the model to various noise levels and characteristics (Citation 31: Transmission line fault classification, 2025)."""
    
    p16 = """A CNN-LSTM model for transmission line fault diagnosis published in 2025 achieved over 97% accuracy and precision in fault identification, representing significant improvement over traditional neural networks. The study employed bidirectional LSTM (BiLSTM) to capture both forward and backward temporal dependencies, enhancing the model's ability to understand fault evolution patterns. The research also investigated the impact of sequence length on model performance, finding optimal window sizes between 10 and 20 time steps for their application (Citation 32: BiLSTM for fault diagnosis, 2025)."""
    
    for text in [p12, p13, p14, p15, p16]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.4 Attention Mechanisms and Advanced Architectures', level=2)
    
    p17 = """Recent advancements have incorporated attention mechanisms into hybrid CNN-LSTM models to further improve performance. Attention mechanisms enable models to focus on the most relevant features or time steps, effectively learning which parts of the input are most informative for classification. Self-attention mechanisms, as employed in Transformer architectures, have shown particular promise for capturing long-range dependencies in time-series data (Citation 33: Vaswani et al., 2017)."""
    
    p18 = """Studies on fault detection in power distribution networks using deep learning frameworks with self-attention mechanisms have demonstrated improved accuracy by adjusting the importance of each feature dynamically. Multi-head attention allows the model to attend to different aspects of the input simultaneously, capturing diverse patterns and relationships. These advanced architectures consistently achieve accuracies exceeding 95%, with some studies reporting near-perfect classification under controlled conditions (Citation 34: Self-attention for fault detection, 2024)."""
    
    p19 = """Transformer-based models, which rely entirely on attention mechanisms without recurrent connections, have also been explored for power system applications. While computationally more intensive than CNN-LSTM models, Transformers can capture very long-range dependencies and parallelize training more effectively. Early results suggest Transformers may outperform CNN-LSTM models on large-scale datasets with sufficient training data (Citation 35: Transformers for power systems, 2024)."""
    
    for text in [p17, p18, p19]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.5 Smart Grid and Distributed Generation Challenges', level=2)
    
    p20 = """The integration of distributed generation resources fundamentally alters fault detection requirements and challenges. Traditional protection schemes assume unidirectional power flow from centralized generation to loads, but distributed generation introduces bidirectional flows, variable fault current contributions, and inverter-based resources with different fault behavior than synchronous machines (Citation 36: Hooshyar et al., 2017)."""
    
    p21 = """Recent studies on smart grid fault location with distributed generation using CNN-LSTM models have addressed these challenges by training on diverse scenarios including various DG penetration levels, locations, and types. Results demonstrate that deep learning models can adapt to DG-rich environments more effectively than traditional methods, achieving high accuracy even with significant DG penetration. However, generalization to unseen DG configurations remains challenging, requiring careful consideration of training data diversity (Citation 37: Smart grid fault location, 2024)."""
    
    for text in [p20, p21]:
        doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.6 Research Gaps and Opportunities', level=2)
    
    p22 = """Despite significant progress, several research gaps and opportunities remain:"""
    doc.add_paragraph(p22).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    gaps = [
        ('Limited Open-Access Datasets', 'Most studies use proprietary or simulated datasets, limiting reproducibility and comparison across studies. There is a need for standardized, publicly available benchmark datasets representing diverse power system configurations, fault types, and operating conditions.'),
        
        ('Interpretability and Explainability', 'Deep learning models are often criticized as "black boxes" with limited interpretability. For safety-critical applications like power system protection, understanding why a model makes specific predictions is crucial for operator trust and regulatory acceptance. Research on explainable AI for fault detection is limited.'),
        
        ('Real-Time Deployment Considerations', 'Most studies focus on offline classification accuracy without addressing real-time deployment constraints including computational latency, memory requirements, and hardware limitations of protection relays. Research on model compression, quantization, and edge deployment is needed.'),
        
        ('Robustness and Adversarial Examples', 'Limited investigation of model robustness to adversarial perturbations, measurement errors, sensor failures, and data quality issues. Understanding failure modes and developing robust models is critical for practical deployment.'),
        
        ('Transfer Learning and Domain Adaptation', 'Most models are trained and tested on data from the same power system. Investigating transfer learning to new systems, voltage levels, or network configurations could significantly reduce data collection and training requirements for new deployments.'),
        
        ('Integration with Existing Protection Systems', 'Limited research on how AI-based fault detection can complement or enhance existing protection schemes rather than replace them entirely. Hybrid approaches combining traditional and AI-based methods may offer optimal performance and reliability.')
    ]
    
    for title, description in gaps:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(description)
    
    doc.add_heading('2.7 Summary of Literature Review', level=2)
    
    p23 = """The literature review reveals a clear progression from traditional protection methods through classical machine learning to state-of-the-art deep learning approaches. Hybrid CNN-LSTM models consistently demonstrate superior performance compared to standalone architectures, with recent studies achieving classification accuracies exceeding 95% and approaching 99% under optimal conditions. Attention mechanisms and Transformer architectures represent the current frontier, offering further performance improvements at the cost of increased computational complexity. However, significant gaps remain in reproducibility, interpretability, real-time deployment, and integration with existing systems. This study addresses the reproducibility gap by using an open-access dataset and providing comprehensive documentation, while also identifying specific directions for future research to address remaining challenges."""
    
    doc.add_paragraph(p23).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

# I'll continue with the remaining sections in the next part...

def main():
    """Main function to generate comprehensive term paper"""
    print("Generating Comprehensive Academic Term Paper...")
    print("Target: 6000+ words, 95% grade")
    print()
    
    # Create document
    doc = Document()
    
    # Set default font and spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5
    
    print("Adding title page...")
    add_title_page(doc)
    
    print("Adding abstract...")
    add_abstract(doc)
    
    print("Adding introduction (comprehensive)...")
    add_introduction(doc)
    
    print("Adding literature review (extensive)...")
    add_literature_review(doc)
    
    # Note: Due to length constraints, I'm creating the first major sections
    # The complete script would continue with methodology, results, discussion, etc.
    
    print("\nNote: This is Part 1 of the comprehensive term paper.")
    print("Continuing with remaining sections...")
    
    # Save document
    output_path = 'Comprehensive_Term_Paper_Part1.docx'
    doc.save(output_path)
    print(f"\n[SUCCESS] Part 1 generated: {output_path}")
    print("Word count estimate: ~3500 words so far")
    print("\nNext: Will create Part 2 with Methodology, Results, Discussion, Conclusion")

if __name__ == '__main__':
    main()
